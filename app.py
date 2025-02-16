# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file
import os
from datetime import datetime, timedelta
import json

app = Flask(__name__)

# Percorso del file JSON per salvare i dati
DATA_FILE = 'tasks.json'

# Funzione per formattare le date (filtro Jinja2)
def strftime_filter(value, format='%d/%m/%Y %H:%M'):
    if value is None:
        return ""  # Gestisci eventuali valori None
    return datetime.strptime(value, '%Y-%m-%d %H:%M:%S').strftime(format)

# Registra il filtro strftime in Jinja2
app.jinja_env.filters['strftime'] = strftime_filter

# Funzione per caricare i dati dal file JSON
def load_tasks():
    if not os.path.exists(DATA_FILE):
        return []
    with open(DATA_FILE, 'r') as f:
        try:
            return json.load(f)
        except json.JSONDecodeError:
            return []

# Funzione per salvare i dati nel file JSON
def save_tasks(tasks):
    with open(DATA_FILE, 'w') as f:
        json.dump(tasks, f, indent=4)

# Route principale per visualizzare la Kanban Board
@app.route('/', methods=['GET', 'POST'])
def kanban():
    status_filter = request.args.get('status', None)  # Filtro per stato
    priority_filter = request.args.get('priority', None)  # Filtro per priorità
    tasks = load_tasks()

    if status_filter and priority_filter:
        tasks = [task for task in tasks if task['status'] == status_filter and task['priority'] == priority_filter]
    elif status_filter:
        tasks = [task for task in tasks if task['status'] == status_filter]
    elif priority_filter:
        tasks = [task for task in tasks if task['priority'] == priority_filter]

    return render_template('kanban.html', tasks=tasks)

# Route per aggiungere una nuova task
@app.route('/add_task', methods=['POST'])
def add_task():
    title = request.form['title']
    priority = request.form.get('priority', 'Medium')  # Priorità predefinita "Medium"
    if not title:
        return redirect(url_for('kanban'))  # Ignora task vuote

    tasks = load_tasks()
    new_task = {
        'id': len(tasks) + 1,
        'title': title,
        'status': 'To Do',
        'priority': priority,
        'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'started_at': None,
        'completed_at': None
    }
    tasks.append(new_task)
    save_tasks(tasks)
    return redirect(url_for('kanban'))

# Route per aggiornare lo stato di una task
@app.route('/update_task/<int:task_id>/<new_status>', methods=['GET'])
def update_task(task_id, new_status):
    tasks = load_tasks()
    for task in tasks:
        if task['id'] == task_id:
            if new_status == 'In Progress':
                task['status'] = 'In Progress'
                task['started_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            elif new_status == 'Done':
                task['status'] = 'Done'
                task['completed_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            else:
                task['status'] = new_status
            break
    save_tasks(tasks)
    return redirect(url_for('kanban'))

# Route per eliminare una task
@app.route('/delete_task/<int:task_id>', methods=['GET'])
def delete_task(task_id):
    tasks = load_tasks()
    tasks = [task for task in tasks if task['id'] != task_id]
    save_tasks(tasks)
    return redirect(url_for('kanban'))

# Route per esportare il report in Word
@app.route('/export_word', methods=['GET'])
def export_word():
    from docx import Document  # Per Word
    tasks = load_tasks()
    report_data = [task for task in tasks if task['status'] == 'Done' and task['completed_at']]

    doc = Document()
    doc.add_heading("Report Settimanale delle Task Completate", level=1)
    if not report_data:
        doc.add_paragraph("Nessuna task completata negli ultimi 7 giorni.")
    else:
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Titolo'
        hdr_cells[1].text = 'Priorità'
        hdr_cells[2].text = 'Creato Il'
        hdr_cells[3].text = 'Iniziato Il'
        hdr_cells[4].text = 'Completato Il'
        for task in report_data:
            row_cells = table.add_row().cells
            row_cells[0].text = task['title']
            row_cells[1].text = task['priority']
            row_cells[2].text = datetime.strptime(task['created_at'], '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y %H:%M')
            row_cells[3].text = datetime.strptime(task['started_at'], '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y %H:%M') if task['started_at'] else "Non applicabile"
            row_cells[4].text = datetime.strptime(task['completed_at'], '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y %H:%M') if task['completed_at'] else "Non applicabile"

    word_filename = "weekly_report.docx"
    doc.save(word_filename)
    return send_file(word_filename, as_attachment=True, download_name="weekly_report.docx")

# Route per esportare la Kanban Board in PDF
@app.route('/export_pdf', methods=['GET'])
def export_pdf():
    import pdfkit  # Per PDF
    tasks = load_tasks()
    html_content = render_template('kanban.html', tasks=tasks)
    pdf_filename = "kanban_board.pdf"
    pdfkit.from_string(html_content, pdf_filename)
    return send_file(pdf_filename, as_attachment=True, download_name="kanban_board.pdf")

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))  # Usa la porta specificata dall'ambiente o la porta 5000
    app.run(host='0.0.0.0', port=port, debug=False)